from flask import Flask, send_file, render_template, request, session, redirect, url_for, jsonify
import time
import json
import os
from math import sqrt
from random import randint, sample, choice as rand_choice
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math2docx

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'

SETTINGS_FILE = Path('data/settings.json')


def load_settings():
    if SETTINGS_FILE.exists():
        with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
            try:
                return json.load(f)
            except json.JSONDecodeError:
                print("Ошибка: Неверный формат JSON в файле настроек.")
                return {
                    'profile_name': 'Гость',
                    'time_limit': False,
                    'task_limit': False,
                    'time_settings': 15,
                    'task_count': 10
                }
    return {
        'profile_name': 'Гость',
        'time_limit': False,
        'task_limit': False,
        'time_settings': 15,
        'task_count': 10
    }

def save_settings(settings):
    try:
        with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False)
    except Exception as e:
        print(f"Ошибка при сохранении настроек: {e}")

def format_time(seconds):
    minutes = int(seconds // 60)
    sec = int(seconds % 60)
    return f"{minutes}:{sec:02d}"

@app.route('/download_result', methods=['POST'])
def download_result():
    profname = request.form.get('profname')
    level_d = request.form.get('level_d')
    mistakes = request.form.get('mistakes')

    print(f"profname: {profname}, level_d: {level_d}, mistakes: {mistakes}")

    try:
        correct_q = int(request.form.get('correct_q'))
        total_q = int(request.form.get('total_q'))
        total_t = int(request.form.get('total_t'))
    except ValueError as e:
        print(f"Ошибка преобразования: {e}")
        return "Некорректные данные", 400

    result_file_path = save_result(profname, correct_q, total_q, level_d, mistakes, total_t)

    return send_file(result_file_path, as_attachment=True)

@app.route('/save_result', methods=['POST'])
def save_result_route():
    # Получение данных из запроса
    profname = request.json.get('profile_name')
    correct_q = request.json.get('correct')
    total_q = request.json.get('total')
    level_d = request.json.get('difficulty')
    mistakes = request.json.get('mistakes')
    total_t = request.json.get('total_time')

    result_file_path = save_result(profname, correct_q, total_q, level_d, mistakes, total_t)

    return send_file(result_file_path, as_attachment=True)

    return jsonify({'status': 'ok'})

@app.route('/check_answer', methods=['POST'])
def check_answer():
    data = request.get_json()
    selected_answer_index = data.get('answer')
    correct_answer = session.get('current_correct_answer')

    selected_answer = session['current_answers'][selected_answer_index]

    session['total'] += 1
    session['current_question'] = session.get('current_question', 0) + 1
    print(f"Выбрано: {selected_answer}, Правильно: {correct_answer} ")

    if selected_answer == correct_answer:
        session['correct'] += 1
    else:
        mistakes = session.get('mistakes', [])
        mistakes.append({
            'question': session['current_equation'],
            'user_answer': selected_answer,
            'correct_answer': correct_answer
        })
        session['mistakes'] = mistakes


    level = session['difficulty']
    if 'stats' not in session:
        session['stats'] = {level: {'correct': 0, 'total': 0}}

    session['stats'][level]['correct'] += 1 if selected_answer == correct_answer else 0
    session['stats'][level]['total'] += 1

    if session.get('task_limit') and session['current_question'] >= session['task_count']:
        return jsonify({'status': 'finished'})

    if session['difficulty'] == 'Легкий':
        current_equation, correct_answer = generate_equation_easy()
    elif session['difficulty'] == 'Средний':
        current_equation, correct_answer = generate_equation_medium()
    else:
        current_equation, correct_answer = generate_equation_hard()

    session['current_equation'] = current_equation
    session['current_correct_answer'] = correct_answer
    session['current_answers'] = generate_choices(correct_answer)

    return jsonify({
        'status': 'ok',
        'current_question': session['current_question'],
        'current_equation': current_equation,
        'answers': session['current_answers']
    })

def generate_numbers_with_perfect_roots(count):
    return sample(range(1, 16), count)

@app.route('/get_current_question')
def get_current_question():
    return jsonify({'current_question': session.get('current_question', 1)})

@app.route('/show_mistakes')
def show_mistakes():
    mistakes = session.get('mistakes', [])
    settings = load_settings()
    return render_template('mistakes.html', mistakes=mistakes, settings=settings)

def timer(seconds, event):
    time.sleep(seconds)
    event.set()

def hard_sqrt():
    return randint(1, 5)


def generate_equation_easy():
    numbers = generate_numbers_with_perfect_roots(randint(1, 2))
    if len(numbers) == 1:
        return rf"\sqrt{ {numbers[0] ** 2} }", numbers[0]
    else:
        operation = rand_choice(['+', '-', '*'])
        if operation == '-' and numbers[0] < numbers[1]:
            numbers[0], numbers[1] = numbers[1], numbers[0]

        a, b = numbers[0] ** 2, numbers[1] ** 2
        if operation == '+':
            answer = numbers[0] + numbers[1]
        elif operation == '-':
            answer = numbers[0] - numbers[1]
        else:
            answer = numbers[0] * numbers[1]
        return rf"\sqrt{ {a} } {operation} \sqrt{ {b} }", answer

def generate_equation_medium():
    while True:
        numbers = generate_numbers_with_perfect_roots(randint(2, 4))
        if len(numbers) == 2:
            a, b = [x**2 for x in numbers]
            answer = a - b
            if answer > 0:
                return rf"(\sqrt{ {a} } + \sqrt{ {b} }) * (\sqrt{ {a} } - \sqrt{ {b} })", answer
        elif len(numbers) == 3:
            a, b, c = [x**2 for x in numbers]
            option = randint(1, 4)
            if option == 1:
                op1 = rand_choice(["+", "*"])
                op2 = rand_choice(["+", "-"])
                if op1 == "*":
                    equation = rf"\sqrt{ {a} } * \sqrt{ {b} } {op2} \sqrt{ {c} }"
                    answer = numbers[0] * numbers[1]
                    answer = answer + numbers[2] if op2 == '+' else answer - numbers[2]
                else:
                    equation = r"\frac{\sqrt{ {" + str(a) + "} } " + op2 + r" \sqrt{ {" + str(b) + r"} }}{\sqrt{ {" + str(c) + r"} }}"
                    numerator = numbers[0] + numbers[1] if op2 == '+' else numbers[0] - numbers[1]
                    answer = numerator // numbers[2]
                if answer > 0:
                    return equation, answer
            elif option == 2:
                equation = rf"\sqrt{ {a} } * (\sqrt{ {b} } + \sqrt{ {c} })"
                answer = numbers[0] * (numbers[1] + numbers[2])
                if answer > 0:
                    return equation, answer
            elif option == 3:
                equation = rf"\sqrt{ {a} }  * \sqrt{ {b} } * \sqrt{ {c} }"
                answer = numbers[0] * numbers[1] * numbers[2]
                if answer > 0:
                    return equation, answer
            elif option == 4:
                equation = rf"(\sqrt{ {a} } + \sqrt{ {b} }) * (\sqrt{ {b} } + \sqrt{ {c} }) * (\sqrt{ {c} } + \sqrt{ {a} })"
                answer = (numbers[0] + numbers[1]) * (numbers[1] + numbers[2]) * (numbers[2] + numbers[0])
                if answer > 0:
                    return equation, answer
        elif len(numbers) == 4:
            option = randint(1, 2)
            a, b, c, d = [x**2 for x in numbers]
            if option == 1:
                op = rand_choice(["+", "-"])
                equation = rf"\sqrt{ {a} } * \sqrt{ {b} } {op} \sqrt{ {c} } * \sqrt{ {d} }"
                part1 = numbers[0] * numbers[1]
                part2 = numbers[2] * numbers[3]
                answer = part1 + part2 if op == '+' else part1 - part2
            else:
                equation = rf"(\sqrt{ {a} } + \sqrt{ {b} } + \sqrt{ {c} } + \sqrt{ {d} }) + |\sqrt{ {a} } - \sqrt{ {b} } - \sqrt{ {c} } - \sqrt{ {d} }|"
                sum_total = sum(numbers)
                abs_diff = abs(numbers[0] - sum(numbers[1:]))
                answer = sum_total + abs_diff
            if answer > 0:
                return equation, answer

def generate_equation_hard():
    while True:
        numbers = generate_numbers_with_perfect_roots(randint(2, 5))
        if len(numbers) == 2:
            a, b = [x**2 for x in numbers]
            answer = int(4 * sqrt(a) * sqrt(b))
            if answer > 0:
                return rf"(\sqrt{ {a} } + \sqrt{ {b} })² - (\sqrt{ {a} } - \sqrt{ {b} })²", answer
        elif len(numbers) == 3:
            a, b, c = [x**2 for x in numbers[:2]] + [numbers[2]]
            equation = r"(\sqrt{" + str(a) + " + " + str(b) + r"} + \sqrt{" + str(c) + r"}) * (\sqrt{" + str(a) + r" + " + str(b) + r"} - " + rf"\sqrt{ {c} })"
            answer = (a + b) - c**2
            if answer > 0:
                return equation, answer
        elif len(numbers) == 4:
            hard = randint(1, 5)
            a, b, c, d = [hard, numbers[0]**2, hard, numbers[1]**2]
            op = rand_choice(["+", "-", "*"])
            equation = rf"{a}\sqrt{ {b} } {op} {c}\sqrt{ {d} }"
            if op == '+':
                answer = a * sqrt(b) + c * sqrt(d)
            elif op == '-':
                answer = a * sqrt(b) - c * sqrt(d)
            else:
                answer = a * c * sqrt(b * d)
            answer = int(round(answer))
            if answer > 0:
                return equation, answer

def generate_choices(correct_answer):
    correct = int(round(correct_answer))
    lower = max(1, correct - 10)
    upper = max(lower + 1, correct + 10)

    choices = [correct]
    while len(choices) < 4:
        new = randint(lower, upper)
        if new not in choices:
            choices.append(new)
    choices = sample(choices, 4)
    return sample(choices[:2], 2) + sample(choices[2:], 2)

def save_result(profname, correct_q, total_q, level_d, mistakes, total_t):

    doc = Document()
    result_file_path = 'Result_of_Gym_MathGym.docx'


    image_path = os.path.join(os.path.dirname(__file__), 'image/icon.png')


    try:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(3.052), height=Inches(0.9945))

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"Ошибка при добавлении изображения: {e}")


    Main_txt = doc.add_heading('Результат тренажёра', level=1)
    Main_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    Main_txt.runs[0].font.name = 'Comic Sans MS'
    Main_txt.runs[0].font.size = Pt(28)
    Main_txt.runs[0].bold = True


    p = doc.add_paragraph(profname)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Comic Sans MS'
    p.runs[0].font.size = Pt(18)

    p = doc.add_paragraph(f'Вы правильно ответили на {correct_q} из {total_q} вопросов')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Comic Sans MS'
    p.runs[0].font.size = Pt(18)

    p = doc.add_paragraph(f'Время выполнения: {format_time(total_t)}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Comic Sans MS'
    p.runs[0].font.size = Pt(18)

    p = doc.add_paragraph(f'Уровень сложности: {level_d}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Comic Sans MS'
    p.runs[0].font.size = Pt(18)

    p = doc.add_paragraph(f'Кол-во ошибок: {total_q - correct_q}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.name = 'Comic Sans MS'
    p.runs[0].font.size = Pt(18)


    doc.add_heading('Ошибки результата:', level=2)

    for mistake in mistakes:
        p = doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run('Вопрос: ').bold = True
        equation = mistake['question']
        math2docx.add_math(p, equation )

        p = doc.add_paragraph()
        p.add_run('Неправильный ответ: ').bold = True
        p.add_run(str(mistake['user_answer']))
        p = doc.add_paragraph()
        p.add_run('Правильный ответ: ').bold = True
        p.add_run(str(mistake['correct_answer']))

    # Сохранение документа
    doc.save(result_file_path)
    return result_file_path


# Flask routes
@app.route('/')
def main():
    settings = load_settings()
    session.update(settings)

    ability_total = 0

    if 'stats' in session:
        # Перебираем уровни сложности
        for level in ['Легкий', 'Средний', 'Сложный']:
            stats = session['stats'].get(level, {'correct': 0, 'total': 0})
            if stats['total'] > 0:
                ability_total += (stats['correct'] / stats['total']) * (1 / 3) * 100

    session['correct'] = 0
    session['total'] = 0
    session['current_question'] = 0
    session['mistakes'] = []

    return render_template('main.html', settings=settings, Ability_total=ability_total)


@app.route('/settings', methods=['GET', 'POST'])
def settings():
    if request.method == 'POST':
        profile_name = request.form.get('profile_name')
        time_limit = 'time_limit' in request.form
        task_limit = 'task_limit' in request.form
        time_settings = request.form.get('time_settings')
        task_count = request.form.get('task_count')

        if time_settings == 'custom':
            custom_time = request.form.get('custom_time')
            time_settings = int(custom_time) if custom_time else 15
        elif time_settings is None:
            time_settings = 15
        else:
            time_settings = int(time_settings)

        if task_count == 'custom':
            custom_tasks = request.form.get('custom_tasks')
            task_count = int(custom_tasks) if custom_tasks else 10
        elif task_count is None:
            task_count = 10
        else:
            task_count = int(task_count)

        settings = {
            "profile_name": profile_name,
            "time_limit": time_limit,
            "task_limit": task_limit,
            "time_settings": time_settings,
            "task_count": task_count
        }
        save_settings(settings)

        session.update(settings)

        return redirect(url_for('settings'))

    settings = load_settings()
    return render_template('settings.html', settings=settings)

@app.route('/choose')
def choose_level():
    return render_template('choose.html')

@app.route('/training/<difficulty>')
def training(difficulty):
    session['difficulty'] = difficulty
    session['start_time'] = time.time()
    session['correct'] = 0
    session['total'] = 0
    session['current_question'] = 0

    settings = load_settings()
    session.update(settings)

    if 'stats' not in session:
        session['stats'] = {}
    if difficulty not in session['stats']:
        session['stats'][difficulty] = {'correct': 0, 'total': 0}

    if difficulty == 'Легкий':
        current_equation, correct_answer = generate_equation_easy()
    elif difficulty == 'Средний':
        current_equation, correct_answer = generate_equation_medium()
    else:
        current_equation, correct_answer = generate_equation_hard()

    session['current_equation'] = current_equation
    session['current_correct_answer'] = correct_answer
    session['current_answers'] = generate_choices(correct_answer)
    print(correct_answer)

    return render_template('training.html', difficulty=difficulty, settings=settings, current_equation=current_equation, answers=session['current_answers'])

@app.route('/result')
def result():
    settings = load_settings()
    total_time = time.time() - session.get('start_time', 0)

    level = session.get('difficulty', 'Легкий')
    stats = session['stats'].get(level, {'correct': 0, 'total': 0})

    session['Ability_easy'] = 0
    session['Ability_medium'] = 0
    session['Ability_hard'] = 0

    for difficulty in ['Легкий', 'Средний', 'Сложный']:
        if difficulty in session['stats']:
            correct = session['stats'][difficulty]['correct']
            total = session['stats'][difficulty]['total']
            if total > 0:
                ability = ((total - correct) / total) * (100 / 3)
                session[f'Ability_{difficulty.lower()}'] = ability
            else:
                session[f'Ability_{difficulty.lower()}'] = 0

    session['Ability_total'] = (session['Ability_easy'] + session['Ability_medium'] + session['Ability_hard'])

    return render_template('result.html',
                           settings=settings,
                           total_time=total_time,
                           correct=session.get('correct', 0),
                           total=session.get('total', 0),
                           difficulty=level,
                           stats=stats)

@app.route('/save-settings', methods=['POST'])
def save_current_settings():
    settings = {
        'profile_name': session.get('profile_name'),
        'time_limit': session.get('time_limit'),
        'task_limit': session.get('task_limit'),
        'time_settings': session.get('time_settings'),
        'task_count': session.get('task_count')
    }
    save_settings(settings)
    return jsonify({'status': 'ok'})

if __name__ == '__main__':
    app.run(host="127.0.0.1", port=5000, debug=True)